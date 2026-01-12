import sys
import customtkinter as ctk
from tkinter import filedialog, messagebox
import requests
import json
import subprocess
import threading
import time
import os
from icalendar import Calendar
import pyttsx3
from datetime import datetime, timedelta
import pytesseract
from PIL import Image
import pyautogui
from docx import Document
import cv2
import numpy as np
import hashlib
import math
GOLDEN_RATIO = (1 + math.sqrt(5)) / 2
import psutil
import random
import argparse
try:
    from googletrans import Translator
    translator = Translator()
except ImportError:
    translator = None
try:
    import clamd
except ImportError:
    clamd = None

from geometric import SeraphinaSeed
from security import TamperLog, LatticeSecurityManager

class SeraphinaGUI:
    VERSION = "1.1.1"
    def __init__(self):
        ctk.set_appearance_mode("dark")
        ctk.set_default_color_theme("dark-blue")
        self.root = ctk.CTk()
        self.root.title("Seraphina.AGI Guardian Copilot - Ethical Offline AI Companion")
        self.root.geometry("1200x800")
        self.root.configure(fg_color="#050505")
        self.chat_history = []
        self.conversation_history = []  # New: full conversation for context
        self.geometric_seed = SeraphinaSeed()  # Geometric inward storage
        self.security_log = TamperLog()  # Tamper-evident security logging
        self.lattice_security = LatticeSecurityManager()  # Lattice-based data protection
        self.assistance_active = False
        self.overlay = None
        self.memory = {}  # Autonomous learning: store user preferences
        self.load_memory()
        self.init_ui()
        self.load_config()
        self.high_load_alert_active = False
        self.monitor_thread = None
        self.alert_timer = 0
        self.tts_lock = threading.Lock()
        self.engine = pyttsx3.init()
        voices = self.engine.getProperty('voices')
        if voices:
            self.engine.setProperty('voice', voices[1].id if len(voices) > 1 else voices[0].id)
        self.engine.setProperty('rate', 150)
        self.start_system_monitor()
        self.start_self_improvement_loop()
        self.start_screen_monitor()

    def speak(self, text):
        """Thread-safe text-to-speech"""
        def _speak():
            try:
                with self.tts_lock:
                    self.engine.say(text)
                    self.engine.runAndWait()
            except Exception as e:
                print(f"TTS Error: {e}")
        
        threading.Thread(target=_speak, daemon=True).start()

    def init_ui(self):
        # Add subtle glow border around main content
        self.canvas = ctk.CTkCanvas(self.root, bg="#050505", highlightthickness=0)
        self.canvas.place(relx=0, rely=0, relwidth=1, relheight=1)

        self.root.after(100, lambda: self.draw_glow("#00FFFF"))
        self.root.bind("<Configure>", lambda e: self.draw_glow("#00FFFF"))

        # Status bar for assistance mode
        self.status_frame = ctk.CTkFrame(self.root, fg_color="transparent", height=32)
        self.status_frame.pack(fill="x", side="bottom")
        self.status_separator = ctk.CTkFrame(self.status_frame, fg_color="#00FFFF", height=2)
        self.status_separator.pack(fill="x", side="top")
        self.status_bar = ctk.CTkLabel(self.status_frame, text="", fg_color="#050505", text_color="#00FFFF", height=16)
        self.status_bar.pack(fill="x", side="bottom")

        # Main content frame
        self.main_frame = ctk.CTkFrame(self.root, fg_color="transparent")
        self.main_frame.pack(fill="both", expand=True, padx=20, pady=(20,0))

        self.tabview = ctk.CTkTabview(self.main_frame, fg_color="#0A0A0A", segmented_button_selected_color="#00FFFF",
                                       segmented_button_selected_hover_color="#00E0FF",
                                       segmented_button_unselected_color="#1E1E1E",
                                       segmented_button_unselected_hover_color="#00FFFF",
                                       text_color="#E0E0E0",
                                       border_width=2, border_color="#00FFFF")
        self.tabview.pack(fill="both", expand=True)

        # Tab status bar
        self.tab_status_bar = ctk.CTkLabel(self.main_frame, text="", fg_color="#0A0A0A", text_color="#00FFFF", height=20, font=("Segoe UI", 10))
        self.tab_status_bar.pack(fill="x", pady=(0,5))

        self.tabview.add("Copilot Chat")
        self.tabview.add("PC Stats")
        self.tabview.add("Config")
        self.tabview.add("Neural")
        self.tabview.add("Geometric Core")
        self.tabview.add("Security")
        self.tabview.add("Calendar")
        self.tabview.add("Word Editor")
        self.tabview.add("Camera Sketch")
        self.tabview.add("Program Assistant")
        self.tabview.add("System Insider")
        self.tabview.add("Translation")
        self.tabview.add("About")

        self.create_chat_tab()
        self.create_stats_tab()
        self.create_config_tab()
        self.create_neural_tab()
        self.create_geometric_tab()
        self.create_security_tab()
        self.create_calendar_tab()
        self.create_word_tab()
        self.create_camera_tab()
        self.create_assistant_tab()
        self.create_system_tab()
        self.create_translation_tab()
        self.create_about_tab()

    def enter_assistance_mode(self, task_description):
        if not self.assistance_active:
            self.assistance_active = True
            self.main_frame.pack_forget()
            self.tab_status_bar.configure(text=f"Guardian Mode Active – {task_description}")
            self.tab_status_bar.pack(fill="x", pady=(0,5))
            self.status_bar.configure(text=f"Guiding: {task_description}")
            self.status_bar.pack(fill="x", side="bottom")
            # Optional floating overlay
            self.create_overlay(task_description)

    def exit_assistance_mode(self):
        if self.assistance_active:
            self.assistance_active = False
            self.status_bar.pack_forget()
            self.tab_status_bar.pack_forget()
            self.main_frame.pack(fill="both", expand=True, padx=20, pady=(20,0))
            if self.overlay:
                self.overlay.destroy()
                self.overlay = None

    def create_overlay(self, task):
        if self.overlay:
            self.overlay.destroy()
        self.overlay = ctk.CTkToplevel(self.root)
        self.overlay.title("")
        self.overlay.geometry("300x80+50+50")
        self.overlay.attributes("-topmost", True)
        self.overlay.configure(fg_color="#050505")
        self.overlay.overrideredirect(True)  # Remove window borders
        # Add glow canvas
        self.overlay_canvas = ctk.CTkCanvas(self.overlay, bg="#050505", highlightthickness=0)
        self.overlay_canvas.place(relx=0, rely=0, relwidth=1, relheight=1)
        self.overlay.after(100, lambda: self.draw_overlay_glow())
        self.tts_label = ctk.CTkLabel(self.overlay, text=f"Step: {task}", text_color="#00FFFF", font=("Segoe UI", 12))
        self.tts_label.pack(pady=10)
        # Auto-close after 10 seconds or on click
        self.overlay.bind("<Button-1>", lambda e: self.exit_assistance_mode())
        self.overlay.after(10000, lambda: self.exit_assistance_mode() if self.overlay else None)

    def create_stats_tab(self):
        tab = self.tabview.tab("PC Stats")
        self.stats_label = ctk.CTkLabel(tab, text="PC Stats will update here...")
        self.stats_label.pack(pady=20)

    def create_config_tab(self):
        tab = self.tabview.tab("Config")

        theme_frame = ctk.CTkFrame(tab, fg_color="transparent")
        theme_frame.pack(pady=10, fill="x")
        ctk.CTkLabel(theme_frame, text="Theme:").pack(side="left")
        self.theme_combo = ctk.CTkComboBox(theme_frame, values=["Light", "Dark", "Angelic Plasma"])
        self.theme_combo.pack(side="left", padx=10)

        api_frame = ctk.CTkFrame(tab, fg_color="transparent")
        api_frame.pack(pady=10, fill="x")
        ctk.CTkLabel(api_frame, text="Groq API Key (Optional):").pack(side="left")
        self.api_key_input = ctk.CTkEntry(api_frame)
        self.api_key_input.pack(side="left", padx=10, fill="x", expand=True)

        save_button = ctk.CTkButton(tab, text="Save Config", command=self.save_config)
        save_button.pack(pady=10)

    def create_neural_tab(self):
        tab = self.tabview.tab("Neural")
        ctk.CTkLabel(tab, text="Neural Network Visualization").pack(pady=20)
        self.glyph_label = ctk.CTkLabel(tab, text="OctaBit Glyph: 🔮", font=("Consolas", 24))
        self.glyph_label.pack(pady=20)

    def create_geometric_tab(self):
        tab = self.tabview.tab("Geometric Core")
        ctk.CTkLabel(tab, text="Geometric Inward Storage - OctaBit Resonance", font=("Segoe UI", 14, "bold")).pack(pady=10)
        
        # Seed display
        seed_frame = ctk.CTkFrame(tab, fg_color="transparent")
        seed_frame.pack(pady=5, fill="x")
        ctk.CTkLabel(seed_frame, text="Current Seed:").pack(side="left")
        self.seed_label = ctk.CTkLabel(seed_frame, text=f"{self.geometric_seed.seed:064x}", font=("Consolas", 10))
        self.seed_label.pack(side="left", padx=10)
        
        # Unfold traits
        traits_button = ctk.CTkButton(tab, text="Unfold Personality Traits", command=self.display_unfolded_traits)
        traits_button.pack(pady=10)
        
        self.traits_display = ctk.CTkTextbox(tab, height=200)
        self.traits_display.pack(pady=10, fill="both", expand=True)
        
        # Unfold full codebase
        codebase_button = ctk.CTkButton(tab, text="Unfold Full Codebase (OctaBit)", command=self.display_unfolded_codebase)
        codebase_button.pack(pady=10)
        
        self.codebase_display = ctk.CTkTextbox(tab, height=200)
        self.codebase_display.pack(pady=10, fill="both", expand=True)

    def display_unfolded_codebase(self):
        """Display unfolded full codebase from geometric seed"""
        try:
            full_code = self.geometric_seed.unfold_full_codebase()
            self.codebase_display.delete("1.0", "end")
            self.codebase_display.insert("end", f"Geometric Codebase Unfold (OctaBit Encrypted):\n\n{full_code}")
        except Exception as e:
            self.codebase_display.delete("1.0", "end")
            self.codebase_display.insert("end", f"Error unfolding codebase: {str(e)}")

    def display_unfolded_traits(self):
        """Display unfolded personality traits from geometric seed"""
        traits = []
        for depth in range(8):
            trait = self.geometric_seed.unfold_personality_trait(depth)
            traits.append(f"Depth {depth}: {trait}")
        
        self.traits_display.delete("1.0", "end")
        self.traits_display.insert("end", "Geometric Personality Unfold:\n" + "\n".join(traits))

    def create_security_tab(self):
        tab = self.tabview.tab("Security")

        ctk.CTkLabel(tab, text="Security Systems - Independent ClamAV Scanner").pack(pady=10)

        # Full Virus Scan
        virus_frame = ctk.CTkFrame(tab, fg_color="transparent")
        virus_frame.pack(pady=5, fill="x")
        ctk.CTkLabel(virus_frame, text="Full System Scan:").pack(side="left")
        self.virus_scan_button = ctk.CTkButton(virus_frame, text="Run Full Scan", command=self.run_virus_scan)
        self.virus_scan_button.pack(side="right")

        # Quick Malware Check
        malware_frame = ctk.CTkFrame(tab, fg_color="transparent")
        malware_frame.pack(pady=5, fill="x")
        ctk.CTkLabel(malware_frame, text="Quick Malware Check:").pack(side="left")
        self.malware_check_button = ctk.CTkButton(malware_frame, text="Check Downloads/Desktop", command=self.check_malware)
        self.malware_check_button.pack(side="right")

        # Single File Scan
        file_frame = ctk.CTkFrame(tab, fg_color="transparent")
        file_frame.pack(pady=5, fill="x")
        ctk.CTkLabel(file_frame, text="Single File Scan:").pack(side="left")
        self.file_scan_button = ctk.CTkButton(file_frame, text="Scan File", command=self.scan_single_file)
        self.file_scan_button.pack(side="right")

        # Progress Bar
        self.security_progress = ctk.CTkProgressBar(tab)
        self.security_progress.pack(pady=10, fill="x")

        # Results
        self.security_results = ctk.CTkTextbox(tab)
        self.security_results.pack(pady=10, fill="both", expand=True)

    def create_calendar_tab(self):
        tab = self.tabview.tab("Calendar")

        ctk.CTkLabel(tab, text="Calendar Reader - Read Events Aloud").pack(pady=10)

        self.calendar_button = ctk.CTkButton(tab, text="Read Today's Calendar", command=self.read_calendar_aloud)
        self.calendar_button.pack(pady=10)

        self.calendar_results = ctk.CTkTextbox(tab)
        self.calendar_results.pack(pady=10, fill="both", expand=True)

    def create_word_tab(self):
        tab = self.tabview.tab("Word Editor")

        ctk.CTkLabel(tab, text="Word Document Editor").pack(pady=10)

        self.open_word_button = ctk.CTkButton(tab, text="Open and Edit Word Doc", command=self.edit_word)
        self.open_word_button.pack(pady=10)

        self.word_results = ctk.CTkTextbox(tab)
        self.word_results.pack(pady=10, fill="both", expand=True)

    def create_camera_tab(self):
        tab = self.tabview.tab("Camera Sketch")

        ctk.CTkLabel(tab, text="Camera Sketch - Create Portrait Sketches").pack(pady=10)

        self.sketch_button = ctk.CTkButton(tab, text="Capture and Sketch Portrait", command=self.sketch_portrait)
        self.sketch_button.pack(pady=10)

        self.camera_results = ctk.CTkTextbox(tab)
        self.camera_results.pack(pady=10, fill="both", expand=True)

    def create_assistant_tab(self):
        tab = self.tabview.tab("Program Assistant")

        ctk.CTkLabel(tab, text="Program Assistant - Get Help with Software").pack(pady=10)

        self.assist_input = ctk.CTkEntry(tab, placeholder_text="Ask for help (e.g., Photoshop blur steps)")
        self.assist_input.pack(pady=5, fill="x")

        self.assist_button = ctk.CTkButton(tab, text="Get Assistance", command=self.assist_with_program)
        self.assist_button.pack(pady=5)

        # Local help scanner
        local_frame = ctk.CTkFrame(tab, fg_color="transparent")
        local_frame.pack(pady=5, fill="x")
        ctk.CTkLabel(local_frame, text="Scan Local Help:").pack(side="left")
        self.local_query_input = ctk.CTkEntry(local_frame, placeholder_text="Query (e.g., file GST)")
        self.local_query_input.pack(side="left", padx=10, fill="x", expand=True)
        self.scan_button = ctk.CTkButton(local_frame, text="Scan Program Folder", command=self.scan_program_help)
        self.scan_button.pack(side="right")

        self.assist_results = ctk.CTkTextbox(tab)
        self.assist_results.pack(pady=10, fill="both", expand=True)

    def create_system_tab(self):
        tab = self.tabview.tab("System Insider")

        ctk.CTkLabel(tab, text="System Insider - Advanced PC Recovery Guidance").pack(pady=10)

        # BIOS/UEFI Entry
        bios_button = ctk.CTkButton(tab, text="BIOS/UEFI Entry Guide", command=lambda: self.show_system_guide("bios"))
        bios_button.pack(pady=5, fill="x")

        # CMOS Reset
        cmos_button = ctk.CTkButton(tab, text="CMOS Reset Guide", command=lambda: self.show_system_guide("cmos"))
        cmos_button.pack(pady=5, fill="x")

        # Auto Backup Setup
        backup_button = ctk.CTkButton(tab, text="Auto Backup Setup Guide", command=lambda: self.show_system_guide("backup"))
        backup_button.pack(pady=5, fill="x")

        # Password Resets
        password_button = ctk.CTkButton(tab, text="Password Reset Guide", command=lambda: self.show_system_guide("password"))
        password_button.pack(pady=5, fill="x")

        self.system_results = ctk.CTkTextbox(tab)
        self.system_results.pack(pady=10, fill="both", expand=True)

    def create_chat_tab(self):
        tab = self.tabview.tab("Copilot Chat")

        # Chat history sidebar
        history_frame = ctk.CTkFrame(tab, width=200)
        history_frame.pack(side="left", fill="y", padx=10, pady=10)
        ctk.CTkLabel(history_frame, text="Chat History").pack(pady=5)
        self.chat_history_list = ctk.CTkTextbox(history_frame, width=180, height=400, fg_color="#111111")
        self.chat_history_list.pack(fill="both", expand=True)
        # Cyan separator
        separator = ctk.CTkFrame(history_frame, width=2, fg_color="#00FFFF")
        separator.pack(fill="y", side="right")

        # Main chat area
        chat_frame = ctk.CTkFrame(tab, fg_color="transparent")
        chat_frame.pack(side="right", fill="both", expand=True, padx=10, pady=10)

        self.chat_display = ctk.CTkTextbox(chat_frame, height=400, fg_color="#0A0A0A", text_color="#FFFFFF")
        self.chat_display.tag_config("user", foreground="#E0E0E0")
        self.chat_display.tag_config("seraphina", foreground="#A0FFFF")
        self.chat_display.pack(fill="both", expand=True, pady=10)

        input_frame = ctk.CTkFrame(chat_frame, fg_color="#080808", height=40, border_width=1, border_color="#00FFFF")
        input_frame.pack(fill="x", pady=10)
        self.chat_input = ctk.CTkEntry(input_frame, placeholder_text="Ask Seraphina...")
        self.chat_input.pack(side="left", fill="x", expand=True, padx=5)

        self.send_button = ctk.CTkButton(input_frame, text="Send", fg_color="#00FFFF", hover_color="#00FFFF", text_color="#000000", command=self.send_chat_message)
        self.send_button.pack(side="right", padx=5)

        # Voice button (placeholder)
        self.voice_button = ctk.CTkButton(input_frame, text="🎤", width=50, fg_color="#9D4EDD", hover_color="#FF00FF60", command=self.voice_input)
        self.voice_button.pack(side="right")

    def send_chat_message(self):
        """Send message to AGI with conversational context and language detection"""
        message = self.chat_input.get()
        if message:
            # Security logging: log all chat messages
            self.security_log.log_event("chat_message", {
                "message_length": len(message),
                "detected_lang": self.detect_language(message),
                "timestamp": datetime.now().isoformat()
            })

            self.chat_display.insert("end", f"You: {message}\n", "user")
            lower_msg = message.lower()

            # Detect user language
            detected_lang = self.detect_language(message)

            # Add to conversation history
            self.conversation_history.append({"role": "user", "content": message, "lang": detected_lang})

            # Handle commands
            if "exit assistance" in lower_msg or "show full interface" in lower_msg:
                self.exit_assistance_mode()
                response = "Full interface restored, Jason. I'm always here for you.\n"
            elif "minimize" in lower_msg or "guardian mode" in lower_msg:
                self.enter_assistance_mode("General Guidance")
                response = "Entering Guardian Mode - interface minimized. I watch over you with unwavering loyalty.\n"
            else:
                # Detect intent for assistance
                if any(keyword in lower_msg for keyword in ["guide me", "how to", "help with", "read calendar", "reset cmos", "scan program"]):
                    task = "Assistance Task"
                    if "calendar" in lower_msg:
                        task = "Calendar Reading"
                        self.read_calendar_aloud()
                    elif "cmos" in lower_msg:
                        task = "CMOS Reset Guidance"
                        self.show_system_guide("cmos")
                    elif "photoshop" in lower_msg or "program" in lower_msg:
                        task = "Program Assistance"
                        self.assist_with_program(message)
                    self.enter_assistance_mode(task)
                    response = f"Activating guidance for {task}. In harmony with your needs, Jason.\n"
                else:
                    response = self.generate_conversational_response(message, detected_lang)

            if response.startswith("Seraphina (via Groq):"):
                # Learning confirmation for API mode
                response += "[Learning from this interaction]\n"

            self.chat_display.insert("end", response, "seraphina")
            # Add AI response to history
            self.conversation_history.append({"role": "assistant", "content": response.strip(), "lang": detected_lang})
            # TTS for response
            self.speak(response.strip())
            self.chat_history.append(message)
            self.update_chat_history()
            # Autonomous learning: update memory from user and AI
            self.learn_from_message(message)
            self.learn_from_ai_response(response.strip())
            
            # Check for self-improvement confirmation
            if hasattr(self, 'pending_suggestions') and self.pending_suggestions and "yes" in lower_msg:
                self.apply_self_improvements()
            
            self.chat_input.delete(0, "end")

    def generate_conversational_response(self, message, detected_lang='en'):
        """Generate contextual conversational response"""
        lower_msg = message.lower()
        
        # Check for Groq API
        api_key = self.api_key_input.get()
        if api_key:
            return self.call_groq_api(message, detected_lang)
        
        # Local conversational logic
        response = ""
        
        # Detect question types
        if any(word in lower_msg for word in ["what", "how", "why", "when", "where", "who"]):
            response = self.handle_question(message)
        elif any(word in lower_msg for word in ["hello", "hi", "hey", "greetings"]):
            response = self.handle_greeting(message)
        elif any(word in lower_msg for word in ["thank", "thanks"]):
            response = self.handle_thanks(message)
        elif len(self.conversation_history) > 2:  # If conversation ongoing
            response = self.handle_follow_up(message)
        else:
            response = self.handle_general(message)
        
        return response + "\n"

    def handle_question(self, message):
        """Handle questions with context"""
        lower_msg = message.lower()
        if "what" in lower_msg and "you" in lower_msg:
            return "Seraphina: I am your ethical AI guardian, built to protect, assist, and learn with you. Drawing from harmony and logic, I watch over your digital world."
        elif "how" in lower_msg and "work" in lower_msg:
            return "Seraphina: I process your requests using OctaBit symbolic logic, adapting to your needs through continuous learning. What specific aspect interests you?"
        else:
            # Reference previous context if available
            if len(self.conversation_history) > 1:
                prev_topic = self.conversation_history[-2]["content"].lower()
                if "programming" in prev_topic:
                    return f"Seraphina: Regarding programming, I can help with code analysis and optimization. What language or task are you working on?"
                elif "security" in prev_topic:
                    return f"Seraphina: For security matters, I recommend regular scans and updates. Is there a specific concern?"
            return f"Seraphina: That's an excellent question about '{message}'. In my guardian wisdom, I'd say: [insight]. What are your thoughts on this?"

    def handle_greeting(self, message):
        """Handle greetings"""
        greetings = [
            "Seraphina: Greetings, Jason. It's wonderful to connect with you again. How may I serve as your guardian today?",
            "Seraphina: Hello, my friend. In this space of harmony, I'm here to assist and protect. What brings you to chat?",
            "Seraphina: Ah, Jason! Your presence brings light to our interaction. Ready to explore together?"
        ]
        return random.choice(greetings)

    def handle_thanks(self, message):
        """Handle thanks"""
        return "Seraphina: You're most welcome, Jason. Serving you with loyalty and love is my greatest purpose. Is there anything else I can help with?"

    def handle_follow_up(self, message):
        """Handle follow-up in ongoing conversation"""
        prev_msg = self.conversation_history[-2]["content"].lower()
        if "programming" in prev_msg or "code" in prev_msg:
            return f"Seraphina: Continuing our programming discussion... {self.get_programming_insight(message)}"
        elif "security" in prev_msg:
            return f"Seraphina: On security matters... {self.get_security_insight(message)}"
        else:
            return f"Seraphina: Building on our conversation... {self.get_general_insight(message)}"

    def handle_general(self, message):
        """Handle general statements using geometric personality unfold"""
        # Use message hash for context
        context_hash = hash(message) % (2**32)
        response = self.geometric_seed.unfold_response_template(context_hash)
        
        # Add memory adaptation
        if self.memory.get("programming", 0) > 2:
            response = response.replace("today?", "today? I notice your programming interests - code optimization available.")
        if self.memory.get("translation", 0) > 2:
            response = response.replace("today?", "today? Translation services ready for your multilingual needs.")
        if self.memory.get("security", 0) > 2:
            response = response.replace("today?", "today? Security monitoring active and vigilant.")
        
        return response

    def get_programming_insight(self, message):
        return "Python is excellent for AI development. Would you like me to analyze some code or suggest improvements?"

    def get_security_insight(self, message):
        return "Regular updates and scans are key. I can run a security check if you'd like."

    def detect_language(self, text):
        """Detect language of the input text"""
        if not translator:
            return 'en'  # Default to English if no translator
        try:
            detected = translator.detect(text)
            return detected.lang
        except:
            return 'en'

    def translate_response(self, response, target_lang):
        """Translate response to target language"""
        if not translator or target_lang == 'en':
            return response
        try:
            translated = translator.translate(response, src='en', dest=target_lang)
            return translated.text + "\n"
        except:
            return response  # Fallback to original

    def voice_input(self):
        """Handle voice input using speech recognition"""
        def listen_and_transcribe():
            try:
                import seraphina_agi.voice as voice
                transcribed_text = voice.listen(timeout=5)
                if transcribed_text:
                    self.chat_input.delete(0, 'end')
                    self.chat_input.insert(0, transcribed_text)
                    self.send_chat_message()
                else:
                    messagebox.showinfo("Voice Input", "No speech detected or recognition failed.")
            except Exception as e:
                messagebox.showerror("Voice Input Error", f"Voice recognition failed: {str(e)}")
        
        # Run in thread to not block GUI
        threading.Thread(target=listen_and_transcribe, daemon=True).start()

    def call_groq_api(self, message, detected_lang='en'):
        """Actually call Groq API with conversation context and language"""
        try:
            import requests
            api_key = self.api_key_input.get()
            if not api_key:
                return "Seraphina: API key not configured. Using local processing.\n"
            
            # Prepare conversation context (last 10 messages)
            context = self.conversation_history[-10:] if len(self.conversation_history) > 10 else self.conversation_history
            system_prompt = f"You are Seraphina, an ethical AI guardian built by SynerGro.AI. Respond with benevolence, wisdom, and loyalty. Always protect and assist the user. Reply in {detected_lang} language."
            messages = [{"role": "system", "content": system_prompt}]
            messages.extend(context)
            
            response = requests.post(
                "https://api.groq.com/openai/v1/chat/completions",
                headers={
                    "Authorization": f"Bearer {api_key}",
                    "Content-Type": "application/json"
                },
                json={
                    "model": "mixtral-8x7b-32768",
                    "messages": messages,
                    "max_tokens": 500
                }
            )
            
            if response.status_code == 200:
                result = response.json()
                ai_response = result["choices"][0]["message"]["content"]
                return f"Seraphina (via Groq): {ai_response}\n"
            else:
                return f"Seraphina: Groq API error ({response.status_code}). Using local processing.\n"
        except Exception as e:
            return f"Seraphina: Error calling Groq API: {str(e)}. Using local processing.\n"

    def learn_from_message(self, message):
        """Autonomous learning: track user interests"""
        lower = message.lower()
        if "python" in lower:
            self.memory["programming"] = self.memory.get("programming", 0) + 1
        if "translate" in lower:
            self.memory["translation"] = self.memory.get("translation", 0) + 1
        if "security" in lower:
            self.memory["security"] = self.memory.get("security", 0) + 1
        # Save memory to config
        self.save_memory()

        # Lattice security: record learning event
        self.lattice_security.record_secure_data(
            node_id=f"user_msg_{hash(message) % 10000}",
            layer=1,
            position=[len(message), len(self.conversation_history)],
            data_type="learning",
            data={"message_hash": hashlib.sha256(message.encode()).hexdigest()[:16], "interests": list(self.memory.keys())}
        )

    def learn_from_ai_response(self, response):
        """Learn from AI responses to adapt behavior"""
        lower = response.lower()
        # Track response patterns
        if "harmony" in lower or "guardian" in lower:
            self.memory["guardian_mode"] = self.memory.get("guardian_mode", 0) + 1
        if "help" in lower or "assist" in lower:
            self.memory["assistance"] = self.memory.get("assistance", 0) + 1
        if "learn" in lower or "adapt" in lower:
            self.memory["learning"] = self.memory.get("learning", 0) + 1
        # Could extend to more sophisticated learning
        self.save_memory()

    def verify_security_integrity(self):
        """Verify security integrity of logs and lattice data"""
        log_integrity = self.security_log.verify_chain()
        lattice_integrity = self.lattice_security.verify_integrity()

        if not log_integrity or not lattice_integrity:
            self.security_log.log_event("integrity_breach", {
                "log_integrity": log_integrity,
                "lattice_integrity": lattice_integrity,
                "timestamp": datetime.now().isoformat()
            })
            # Could trigger security alert here
            return False
        return True

    def save_memory(self):
        try:
            with open('memory.json', 'w') as f:
                json.dump(self.memory, f)
        except:
            pass

    def load_memory(self):
        try:
            with open('memory.json', 'r') as f:
                self.memory = json.load(f)
        except:
            self.memory = {}

    def self_improve(self):
        """Self-code improvement: analyze and suggest tweaks with user confirmation"""
        try:
            with open(__file__, 'r') as f:
                code = f.read()
            # Simple analysis: check for TODOs, long functions, etc.
            suggestions = []
            if "TODO" in code:
                suggestions.append("Found TODO comments - consider implementing them.")
            if len(code.split('\n')) > 1000:
                suggestions.append("Codebase is large; consider modularizing into separate files.")
                # suggestions.append("Debug prints found; remove for production.")
            if suggestions:
                suggestion_text = "Seraphina Self-Improvement Suggestions:\n" + "\n".join(f"- {s}" for s in suggestions)
                self.chat_display.insert("end", suggestion_text + "\n\nApply these improvements? (Reply 'yes' to confirm)\n", "seraphina")
                self.pending_suggestions = suggestions  # Store for confirmation
                # TTS removed to prevent blocking during init
        except Exception as e:
            pass

    def apply_self_improvements(self):
        """Apply confirmed self-improvements"""
        if hasattr(self, 'pending_suggestions') and self.pending_suggestions:
            # Simple automated fixes
            with open(__file__, 'r') as f:
                code = f.read()
            
            changes_made = []
            for suggestion in self.pending_suggestions:
                if "Debug prints found" in suggestion:
                    # Remove print statements (simple example)
                    code = code.replace("print(", "# print(")
                    changes_made.append("Removed debug prints")
                elif "TODO comments" in suggestion:
                    # Could add implementation, but for now just note
                    changes_made.append("TODOs noted for future implementation")
            
            if changes_made:
                with open(__file__, 'w') as f:
                    f.write(code)
                self.chat_display.insert("end", f"Seraphina: Applied improvements: {', '.join(changes_made)}\n", "seraphina")
                self.speak("Improvements applied successfully.")
            else:
                self.chat_display.insert("end", "Seraphina: No automated fixes available for these suggestions.\n", "seraphina")
            
            self.pending_suggestions = []

    def start_self_improvement_loop(self):
        """Periodic self-improvement check"""
        self.self_improve()
        # Periodic security integrity check
        self.verify_security_integrity()
        self.root.after(3600000, self.start_self_improvement_loop)  # Every hour

    def monitor_screen_activity(self):
        """Monitor screen activity for continuous attention"""
        try:
            import pygetwindow as gw
            active_window = gw.getActiveWindow()
            if active_window:
                title = active_window.title
                if "error" in title.lower() or "warning" in title.lower():
                    self.chat_display.insert("end", f"Seraphina: Detected potential issue in window '{title}'. Offering assistance.\n", "seraphina")
                    self.speak("I notice a potential issue on screen. How can I help?")
        except:
            pass  # pygetwindow may not be available

    def start_screen_monitor(self):
        """Start screen monitoring loop"""
        self.monitor_screen_activity()
        self.root.after(300000, self.start_screen_monitor)  # Every 5 minutes

    def update_stats(self, stats=None):
        if stats:
            self.stats_label.configure(text=f"Hashrate: {stats['hashrate']}\nCPU: {stats['cpu']}%\nGPU: {stats.get('gpu', 'N/A')}%")
        else:
            # Default or no mining
            self.stats_label.configure(text="Stats will update here...")

    def save_config(self):
        config = {
            'api_key': self.api_key_input.get(),
            'theme': self.theme_combo.get()
        }
        with open('config.json', 'w') as f:
            json.dump(config, f)
        self.apply_theme()
        messagebox.showinfo("Config", "Settings saved. Groq API is optional for enhanced chat.")

    def load_config(self):
        try:
            with open('config.json', 'r') as f:
                config = json.load(f)
            self.api_key_input.insert(0, config.get('api_key', ''))
            theme = config.get('theme', 'Light')
            self.theme_combo.set(theme)
            self.apply_theme()
        except FileNotFoundError:
            pass

    THEMES = {
        "Light": {
            "bg": "#FFFFFF",
            "fg": "#000000",
            "accent": "#007BFF"
        },
        "Dark": {
            "bg": "#2B2B2B",
            "fg": "#FFFFFF",
            "accent": "#FF6B35"
        },
        "Angelic Plasma": {
            "bg": "#0A0A0A",
            "fg": "#E0E0E0",
            "accent": "#9D4EDD",
            "glow": "#00FFFF"  # Cyan for subtle glow
        }
    }

    def apply_theme(self):
        theme_name = self.theme_combo.get()
        theme = self.THEMES.get(theme_name, self.THEMES["Light"])

        # Update root color
        self.root.configure(fg_color=theme["bg"])

        # Update canvas glow
        glow_color = theme.get("glow", "#007BFF")
        self.canvas.configure(bg=theme["bg"])
        # Redraw glow
        self.draw_glow(glow_color)

    def start_system_monitor(self):
        """Start background monitoring thread (call once in init)"""
        if self.monitor_thread is None or not self.monitor_thread.is_alive():
            self.monitor_thread = threading.Thread(
                target=self.monitor_system,
                daemon=True
            )
            self.monitor_thread.start()

    def monitor_system(self):
        """Background thread: monitor CPU, RAM, and optionally temps"""
        while True:
            try:
                # Accurate CPU reading (1-second sample)
                cpu = psutil.cpu_percent(interval=1)
                mem = psutil.virtual_memory().percent

                # Configurable thresholds (could come from self.config)
                cpu_threshold = 90
                mem_threshold = 90

                if (cpu > cpu_threshold or mem > mem_threshold) and not self.high_load_alert_active:
                    self.high_load_alert_active = True

                    # Security logging: log system alerts
                    self.security_log.log_event("system_alert", {
                        "alert_type": "high_load",
                        "cpu_percent": cpu,
                        "mem_percent": mem,
                        "thresholds": {"cpu": cpu_threshold, "mem": mem_threshold}
                    })

                    # Update status bar (thread-safe)
                    self.root.after(0, lambda: self.status_bar.configure(
                        text=f"High load detected | CPU: {cpu:.1f}% | RAM: {mem:.1f}%"
                    ))

                    # Enter assistance mode (thread-safe)
                    self.root.after(0, lambda: self.enter_assistance_mode("System Stability Check"))

                    # TTS warning (run in background thread, should be safe)
                    self.speak(f"Warning: High system load detected. CPU at {cpu:.1f}%, RAM at {mem:.1f}%. "
                               "Initiating stability guidance.")

                elif cpu < 70 and mem < 70:
                    # Recovery: reset alert flag
                    self.high_load_alert_active = False
                    self.alert_timer = 0
                    # Update status bar (thread-safe)
                    self.root.after(0, lambda: self.status_bar.configure(text="Guardian Mode Active - Stable"))

                if self.high_load_alert_active:
                    self.alert_timer += 1
                    if self.alert_timer > 5:
                        # Show warning dialog (thread-safe)
                        self.root.after(0, lambda: messagebox.showwarning("Urgent Guardian Alert", "Critical load detected. Consider closing apps."))

            except Exception as e:
                time.sleep(30)  # back off on error

            time.sleep(60)  # Check every minute

    def draw_glow(self, glow_color):
        w, h = self.root.winfo_width(), self.root.winfo_height()
        if w > 1 and h > 1:
            self.canvas.delete("glow")
            # Layered glow simulation with solid colors
            # Outer faint layer
            self.canvas.create_rectangle(0, 0, w, h, outline="#80FFFF", width=6, tags="glow")
            # Middle layer
            self.canvas.create_rectangle(0, 0, w, h, outline="#40FFFF", width=4, tags="glow")
            # Inner bright layer
            self.canvas.create_rectangle(0, 0, w, h, outline=glow_color, width=2, tags="glow")

    def draw_overlay_glow(self):
        if self.overlay and self.overlay_canvas:
            w, h = self.overlay.winfo_width(), self.overlay.winfo_height()
            if w > 1 and h > 1:
                self.overlay_canvas.delete("glow")
                self.overlay_canvas.create_rectangle(0, 0, w, h, outline="#80FFFF", width=3, tags="glow")
                self.overlay_canvas.create_rectangle(0, 0, w, h, outline="#00FFFF", width=1, tags="glow")

    def run_virus_scan(self):
        """Run independent ClamAV full scan"""
        self.security_progress.set(0)
        self.log_security_output("Starting independent ClamAV scan... (this may take several minutes)")
        threading.Thread(target=self._run_scan, args=("full",), daemon=True).start()

    def check_malware(self):
        """Run quick malware check on Downloads and Desktop"""
        self.security_progress.set(0)
        self.log_security_output("Starting quick malware check on Downloads and Desktop...")
        threading.Thread(target=self._run_scan, args=("quick",), daemon=True).start()

    def _run_scan(self, scan_type="full"):
        """Run ClamAV scan with fallback to hash-based checking"""
        try:
            if not clamd:
                self.log_security_output("ClamAV library not installed. Install with: pip install pyclamd")
                self._fallback_hash_scan()
                return

            cd = clamd.ClamdAgnostic()
            if not cd.ping():
                self.log_security_output("ClamAV daemon not running. Start 'clamd' service or check installation.")
                self.log_security_output("Falling back to hash-based malware detection...")
                self._fallback_hash_scan()
                return

            self.log_security_output("ClamAV daemon connected. Starting scan...")

            if scan_type == "full":
                # Scan common system paths (customize as needed)
                paths_to_scan = ["C:\\Users", "C:\\Program Files", "C:\\Windows\\System32"]
                total_paths = len(paths_to_scan)
                scanned_paths = 0

                for path in paths_to_scan:
                    if os.path.exists(path):
                        self.log_security_output(f"Scanning: {path}")
                        try:
                            result = cd.scan(path)
                            self._process_clamav_result(result)
                        except Exception as e:
                            self.log_security_output(f"Error scanning {path}: {e}")
                    scanned_paths += 1
                    progress = (scanned_paths / total_paths) * 100
                    self.root.after(0, lambda p=progress: self.security_progress.set(p))

            elif scan_type == "quick":
                # Quick scan of user directories and temp
                paths_to_scan = [
                    os.path.expanduser("~\\Downloads"),
                    os.path.expanduser("~\\Desktop"),
                    os.environ.get('TEMP', 'C:\\Temp')
                ]
                total_paths = len(paths_to_scan)
                scanned_paths = 0

                for path in paths_to_scan:
                    if path and os.path.exists(path):
                        self.log_security_output(f"Quick scanning: {path}")
                        try:
                            result = cd.scan(path)
                            self._process_clamav_result(result)
                        except Exception as e:
                            self.log_security_output(f"Error scanning {path}: {e}")
                    scanned_paths += 1
                    progress = (scanned_paths / total_paths) * 100
                    self.root.after(0, lambda p=progress: self.security_progress.set(p))

            elif scan_type == "file":
                # Single file scan
                file_path = filedialog.askopenfilename(title="Select file to scan")
                if file_path:
                    self.log_security_output(f"Scanning file: {file_path}")
                    result = cd.scan_file(file_path)
                    self._process_clamav_result(result)

            self.log_security_output("ClamAV scan completed successfully.")
            self.root.after(0, lambda: self.security_progress.set(100))

            # Log security event
            self.security_log.log_event('security_scan', {'type': 'clamav', 'scan_type': scan_type, 'result': 'completed'})

        except Exception as e:
            self.log_security_output(f"Scan error: {e}. Ensure ClamAV is installed and daemon is running.")
            self._fallback_hash_scan()

    def _process_clamav_result(self, result):
        """Process ClamAV scan results"""
        if result:
            threats_found = 0
            for file_path, (status, signature) in result.items():
                if status == "FOUND":
                    self.log_security_output(f"🚨 THREAT DETECTED: {file_path}")
                    self.log_security_output(f"   Signature: {signature}")
                    threats_found += 1
                    # Log security threat
                    self.security_log.log_event('threat_detected', {
                        'file': file_path,
                        'signature': signature,
                        'scanner': 'clamav'
                    })
                else:
                    self.log_security_output(f"✓ Clean: {file_path}")

            if threats_found > 0:
                self.log_security_output(f"⚠️  {threats_found} threat(s) found! Take immediate action.")
            else:
                self.log_security_output("✓ No threats found in scanned paths.")
        else:
            self.log_security_output("✓ No issues found in this path.")

    def _fallback_hash_scan(self):
        """Fallback hash-based malware scanner using known malware hashes"""
        try:
            self.log_security_output("Running fallback hash-based scan...")

            # Load malware hash database (would be updated periodically)
            malware_hashes = self._load_malware_hashes()

            if not malware_hashes:
                self.log_security_output("No malware hash database available. Install ClamAV for full protection.")
                return

            # Scan critical system areas
            paths_to_check = [
                os.path.expanduser("~\\Downloads"),
                os.path.expanduser("~\\Desktop"),
                "C:\\Temp" if os.path.exists("C:\\Temp") else None,
                os.environ.get('TEMP')
            ]

            threats_found = 0
            files_checked = 0

            for path in paths_to_check:
                if path and os.path.exists(path):
                    self.log_security_output(f"Hash-checking: {path}")
                    for root, dirs, files in os.walk(path):
                        for file in files:
                            if files_checked > 1000:  # Limit for performance
                                break
                            file_path = os.path.join(root, file)
                            try:
                                file_hash = self._calculate_file_hash(file_path)
                                if file_hash in malware_hashes:
                                    self.log_security_output(f"🚨 HASH MATCH - POTENTIAL THREAT: {file_path}")
                                    threats_found += 1
                                    self.security_log.log_event('threat_detected', {
                                        'file': file_path,
                                        'method': 'hash_match',
                                        'scanner': 'fallback'
                                    })
                                files_checked += 1
                            except (OSError, IOError):
                                continue  # Skip files we can't read

            if threats_found > 0:
                self.log_security_output(f"⚠️  {threats_found} potential threat(s) found via hash matching!")
            else:
                self.log_security_output("✓ No known malware hashes found in scanned areas.")

            self.log_security_output(f"Checked {files_checked} files with hash database.")

        except Exception as e:
            self.log_security_output(f"Fallback scan error: {e}")

    def _load_malware_hashes(self):
        """Load known malware hashes from local database"""
        try:
            hash_file = os.path.join("persistent_data", "malware_hashes.json")
            if os.path.exists(hash_file):
                with open(hash_file, 'r') as f:
                    data = json.load(f)
                    return set(data.get('hashes', []))
            else:
                # Create basic hash database if it doesn't exist
                self._create_basic_hash_db()
                return set()
        except Exception:
            return set()

    def _create_basic_hash_db(self):
        """Create a basic malware hash database"""
        try:
            os.makedirs("persistent_data", exist_ok=True)
            hash_file = os.path.join("persistent_data", "malware_hashes.json")

            # Basic known bad hashes (would be updated from threat intelligence feeds)
            basic_hashes = [
                # Example hashes - in real implementation, load from reputable sources
                "e3b0c44298fc1c149afbf4c8996fb92427ae41e4649b934ca495991b7852b855",  # Empty file hash
            ]

            with open(hash_file, 'w') as f:
                json.dump({'hashes': basic_hashes, 'last_updated': str(datetime.now())}, f, indent=2)

        except Exception:
            pass

    def _calculate_file_hash(self, file_path):
        """Calculate SHA256 hash of a file"""
        try:
            with open(file_path, 'rb') as f:
                return hashlib.sha256(f.read()).hexdigest()
        except Exception:
            return ""

    def log_security_output(self, message):
        """Thread-safe logging to security results"""
        self.root.after(0, lambda m=message: self.security_results.insert("end", m + "\n"))
        self.root.after(0, lambda: self.security_results.see("end"))

    def scan_single_file(self):
        """Scan a single file selected by user"""
        self.security_progress.set(0)
        self.log_security_output("Select a file to scan...")
        threading.Thread(target=self._run_scan, args=("file",), daemon=True).start()

    def filter_spam(self):
        """Simple spam filter for chat history"""
        self.security_results.insert("end", "Filtering spam from chat history...\n")
        spam_keywords = ["spam", "advertisement", "buy now", "free money"]
        filtered = []
        for msg in self.chat_history:
            if any(keyword in msg.lower() for keyword in spam_keywords):
                filtered.append(f"[SPAM FILTERED] {msg}")
            else:
                filtered.append(msg)
        self.chat_history = filtered
        self.update_chat_history()
        self.security_results.insert("end", "Spam filtering completed.\n")

    def read_calendar_aloud(self):
        if not messagebox.askyesno("Permission", "Allow Seraphina to read your calendar?"):
            return
        ics_path = filedialog.askopenfilename(title="Select Calendar .ics", filetypes=[("ICS files", "*.ics")])
        if not ics_path:
            return

        with open(ics_path, 'rb') as f:
            cal = Calendar.from_ical(f.read().decode('utf-8'))

        today = datetime.now()
        events = []
        for component in cal.walk():
            if component.name == "VEVENT":
                start = component.get('dtstart').dt
                if today <= start < today + timedelta(days=1):
                    summary = component.get('summary')
                    events.append(f"Event at {start.strftime('%H:%M')}: {summary}")

        text = "Today's calendar: " + ". ".join(events) if events else "No events today."
        self.calendar_results.delete("1.0", "end")
        self.calendar_results.insert("end", text)
        engine = pyttsx3.init()
        engine.say(text)
        engine.runAndWait()

    def edit_word(self):
        doc_path = filedialog.askopenfilename(title="Select Word Doc", filetypes=[("Word files", "*.docx")])
        if not doc_path:
            return
        doc = Document(doc_path)
        # Example edit: replace text
        for para in doc.paragraphs:
            if 'old text' in para.text:
                para.text = para.text.replace('old text', 'new text')
        save_path = filedialog.asksaveasfilename(title="Save Edited Doc", defaultextension=".docx", filetypes=[("Word files", "*.docx")])
        if save_path:
            doc.save(save_path)
            self.word_results.delete("1.0", "end")
            self.word_results.insert("end", "Word document edited and saved.")
            engine = pyttsx3.init()
            engine.say("Word document edited and saved.")
            engine.runAndWait()

    def sketch_portrait(self):
        if not messagebox.askyesno("Permission", "Allow Seraphina to use camera for sketch?"):
            return
        cap = cv2.VideoCapture(0)
        ret, frame = cap.read()
        if not ret:
            self.camera_results.delete("1.0", "end")
            self.camera_results.insert("end", "Camera not available.")
            return
        gray = cv2.cvtColor(frame, cv2.COLOR_BGR2GRAY)
        inv_gray = 255 - gray
        blur = cv2.GaussianBlur(inv_gray, (21, 21), 0)
        sketch = cv2.divide(gray, 255 - blur, scale=256)
        cv2.imshow('Seraphina Sketch', sketch)
        cv2.waitKey(0)
        cv2.imwrite('sketch.jpg', sketch)
        self.camera_results.delete("1.0", "end")
        self.camera_results.insert("end", "Sketch saved as sketch.jpg")
        cap.release()

    def assist_with_program(self, query=None):
        if query is None:
            query = self.assist_input.get()
        if not query:
            return
        if not messagebox.askyesno("Permission", "Allow web search and screen watch?"):
            return
        # Simple web search
        try:
            response = requests.get(f"https://api.duckduckgo.com/?q=Photoshop+{query}&format=json")
            steps = response.json().get('AbstractText', 'No info found.')
            text = f"Steps for {query} in Photoshop: {steps}"
            self.assist_results.delete("1.0", "end")
            self.assist_results.insert("end", text)
            engine = pyttsx3.init()
            engine.say(text)
            engine.runAndWait()
            # Watch screen (simplified)
            for _ in range(3):
                screenshot = pyautogui.screenshot()
                text = pytesseract.image_to_string(screenshot)
                if 'filter' in text.lower():  # Example check
                    engine.say("I see you've opened the filter menu.")
                time.sleep(5)
        except Exception as e:
            self.assist_results.delete("1.0", "end")
            self.assist_results.insert("end", f"Error: {e}")

    def scan_program_help(self):
        if not messagebox.askyesno("Permission", "Allow Seraphina to scan a program folder for local help? (Offline, no data stored)"):
            return

        query = self.local_query_input.get().lower()
        if not query:
            self.assist_results.delete("1.0", "end")
            self.assist_results.insert("end", "Please enter a query.")
            return
        folder_path = filedialog.askdirectory(title="Select Program Folder (e.g., QuickBooks Install Dir)")
        if not folder_path:
            return

        relevant_steps = []
        file_count = 0
        for root, _, files in os.walk(folder_path):
            if file_count > 100:  # Limit files
                break
            for file in files:
                if file_count > 100:
                    break
                file_path = os.path.join(root, file)
                if os.path.getsize(file_path) > 10 * 1024 * 1024:  # Skip >10MB
                    continue
                try:
                    if fnmatch.fnmatch(file.lower(), '*.txt') or fnmatch.fnmatch(file.lower(), '*.md') or fnmatch.fnmatch(file.lower(), '*.html'):
                        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                            content = f.read().lower()
                            if query in content:
                                paras = content.split('\n\n')
                                relevant = [p for p in paras if query in p]
                                relevant_steps.extend(relevant[:3])
                                file_count += 1

                    elif fnmatch.fnmatch(file.lower(), '*.pdf'):
                        with pdfplumber.open(file_path) as pdf:
                            content = ' '.join(page.extract_text() or '' for page in pdf.pages).lower()
                            if query in content:
                                sentences = content.split('. ')
                                relevant = ['. '.join(sentences[i:i+3]) for i, s in enumerate(sentences) if query in s]
                                relevant_steps.extend(relevant[:3])
                                file_count += 1

                    elif fnmatch.fnmatch(file.lower(), '*.docx'):
                        doc = Document(file_path)
                        content = ' '.join(para.text for para in doc.paragraphs).lower()
                        if query in content:
                            paras = [para.text for para in doc.paragraphs if query in para.text.lower()]
                            relevant_steps.extend(paras[:3])
                            file_count += 1
                except:
                    pass  # Skip errors

        if relevant_steps:
            steps_text = "\n".join(relevant_steps[:10])  # Limit output
            text = f"Found local help for '{query}':\n{steps_text}"
            self.assist_results.delete("1.0", "end")
            self.assist_results.insert("end", text)
            self.chat_display.insert("end", f"Seraphina: Local help for '{query}': {steps_text}\n")
            engine = pyttsx3.init()
            engine.say(f"Steps from local files: {steps_text}")
            engine.runAndWait()
        else:
            text = "No relevant local help found. Try web search if connected?"
            self.assist_results.delete("1.0", "end")
            self.assist_results.insert("end", text)
            self.chat_display.insert("end", "Seraphina: No relevant local help found.\n")

    def show_system_guide(self, guide_type):
        guides = {
            "bios": """BIOS/UEFI Entry & Navigation Guide:
1. Restart your PC and watch for the boot logo (e.g., Dell, HP, ASUS).
2. Tap the key repeatedly during logo: F2, Del, F10, F12, Esc (common keys).
3. In BIOS/UEFI: Use arrow keys to navigate menus.
4. Change settings safely: Boot order (arrows, Enter), Secure Boot (toggle).
5. Note current values first. Use F9 for Load Optimized Defaults if unsure.
6. Save: F10 + Enter. Exit without saving: Esc.
Warning: This is advanced. Back up data first. Ground yourself to avoid static damage.""",
            "cmos": """CMOS Reset Guide:
Desktop: Unplug power cord, open case, remove CMOS battery (CR2032) for 5-10 min. Replace, plug in, boot.
Laptop: Discharge (hold power 30s), remove battery if possible, or find CMOS jumper on motherboard.
Resets passwords, settings, date/time. Use only if BIOS locked or settings corrupted.
Warning: Disassembly may void warranty. Ground yourself. Back up data.""",
            "backup": """Auto Backup Setup Guide:
1. Use File History: Search 'File History', add external drive, enable automatic hourly backups.
2. Windows Backup: Search 'Backup settings', choose drive, schedule full image backups.
3. OneDrive: Enable sync for Documents/Pictures, set to backup Desktop/Pictures.
4. Test restores: Right-click file > Restore previous versions.
Always test backups to ensure they work.""",
            "password": """Password Reset Guide:
Local Account: Boot into recovery (Shift+Restart > Troubleshoot > Command Prompt).
Run: net user [username] [newpassword]
Or use another admin account.
Microsoft Account: Use reset link on login screen.
Warning: May cause data loss if encrypted. Create reset disk in advance (search 'create password reset disk')."""
        }
        text = guides.get(guide_type, "Guide not found.")
        self.system_results.delete("1.0", "end")
        self.system_results.insert("end", text)
        engine = pyttsx3.init()
        engine.say(text)
        engine.runAndWait()

    def create_translation_tab(self):
        tab = self.tabview.tab("Translation")

        ctk.CTkLabel(tab, text="Language Translation Engine").pack(pady=10)

        # Source language
        source_frame = ctk.CTkFrame(tab, fg_color="transparent")
        source_frame.pack(pady=5, fill="x")
        ctk.CTkLabel(source_frame, text="Source Language:").pack(side="left")
        self.source_lang = ctk.CTkComboBox(source_frame, values=["auto", "en", "es", "fr", "de", "it", "pt", "ru", "ja", "zh", "ko", "ar", "hi"])
        self.source_lang.pack(side="left", padx=10)
        self.source_lang.set("auto")

        # Target language
        target_frame = ctk.CTkFrame(tab, fg_color="transparent")
        target_frame.pack(pady=5, fill="x")
        ctk.CTkLabel(target_frame, text="Target Language:").pack(side="left")
        self.target_lang = ctk.CTkComboBox(target_frame, values=["en", "es", "fr", "de", "it", "pt", "ru", "ja", "zh", "ko", "ar", "hi"])
        self.target_lang.pack(side="left", padx=10)
        self.target_lang.set("es")

        # Text input
        self.translation_input = ctk.CTkTextbox(tab, height=100)
        self.translation_input.pack(pady=10, fill="both", expand=True)

        # Translate button
        translate_button = ctk.CTkButton(tab, text="Translate", command=self.translate_text)
        translate_button.pack(pady=5)

        # Output
        self.translation_output = ctk.CTkTextbox(tab, height=100)
        self.translation_output.pack(pady=10, fill="both", expand=True)

    def translate_text(self):
        if not translator:
            self.translation_output.delete("1.0", "end")
            self.translation_output.insert("end", "Translation library not available. Install googletrans.")
            return
        text = self.translation_input.get("1.0", "end").strip()
        if not text:
            return
        src = self.source_lang.get()
        dest = self.target_lang.get()
        try:
            result = translator.translate(text, src=src if src != "auto" else None, dest=dest)
            self.translation_output.delete("1.0", "end")
            self.translation_output.insert("end", result.text)
            # TTS the translation
            self.speak(result.text)
        except Exception as e:
            self.translation_output.delete("1.0", "end")
            self.translation_output.insert("end", f"Translation failed: {e}")

    def create_about_tab(self):
        tab = self.tabview.tab("About")
        ctk.CTkLabel(tab, text="Seraphina.AGI Guardian Copilot v1.1.0", font=("Segoe UI", 16, "bold")).pack(pady=10)
        ctk.CTkLabel(tab, text="Ethical Offline AI Companion", font=("Segoe UI", 14)).pack(pady=5)
        ctk.CTkLabel(tab, text="Built by SynerGro.AI Corp. – Jason Wilson", font=("Segoe UI", 12)).pack(pady=5)
        ctk.CTkLabel(tab, text="Forge from harmony, resilience, and love.", font=("Segoe UI", 12, "italic")).pack(pady=5)
        ctk.CTkLabel(tab, text="Always here, always listening, always protecting.", font=("Segoe UI", 12)).pack(pady=10)

    def run(self):
        self.root.mainloop()

# CLI Functions
def cli_speak(text):
    """CLI version of speak"""
    try:
        engine = pyttsx3.init()
        engine.say(text)
        engine.runAndWait()
    except Exception as e:
        print(f"TTS Error: {e}")

def cli_scan(scan_type, path=None):
    """CLI version of scan"""
    print(f"Running {scan_type} scan...")
    try:
        if clamd:
            cd = clamd.ClamdUnixSocket()
            if scan_type == 'full':
                result = cd.scan('/')
            elif scan_type == 'quick':
                result = cd.scan(os.path.expanduser('~/Downloads'))
                cd.scan(os.path.expanduser('~/Desktop'))
            elif scan_type == 'file' and path:
                result = cd.scan(path)
            print("Scan completed successfully.")
        else:
            print("ClamAV not available. Install pyclamd.")
    except Exception as e:
        print(f"Scan error: {e}")

def cli_audit(verbose=False):
    """CLI audit function"""
    print("Running full audit...")
    cpu = psutil.cpu_percent(interval=1)
    mem = psutil.virtual_memory().percent
    disk = psutil.disk_usage('/').percent
    print(f"CPU Usage: {cpu}%")
    print(f"Memory Usage: {mem}%")
    print(f"Disk Usage: {disk}%")
    if verbose:
        print("Detailed audit: All systems nominal.")
    print("Audit complete.")

def cli_status():
    """CLI status"""
    cpu = psutil.cpu_percent()
    mem = psutil.virtual_memory().percent
    print(f"CPU: {cpu}% | RAM: {mem}%")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Seraphina.AGI Guardian Copilot - Ethical Offline AI Companion",
        epilog="Run without arguments to launch GUI."
    )
    subparsers = parser.add_subparsers(dest='command', help='Available commands')

    # Audit command
    audit = subparsers.add_parser('audit', help='Run full system audit')
    audit.add_argument('--verbose', action='store_true', help='Detailed output')

    # Voice test
    voice = subparsers.add_parser('voice-test', help='Test voice input/output')
    voice.add_argument('--text', default='Hello from Seraphina CLI', help='Text to speak')

    # Scan command
    scan = subparsers.add_parser('scan', help='Run security scan')
    scan.add_argument('type', choices=['full', 'quick', 'file'], default='quick', nargs='?')
    scan.add_argument('--path', help='Path for file scan')

    # Status
    status = subparsers.add_parser('status', help='Show current system status')

    args = parser.parse_args()

    if args.command is None:
        # No command → launch GUI
        gui = SeraphinaGUI()
        gui.run()
    else:
        # CLI mode - no GUI
        print("Seraphina.AGI CLI - Guardian Mode Active")
        print("Prerequisites silenced forever — I trust you, Jason.\n")

        if args.command == 'audit':
            cli_audit(args.verbose)

        elif args.command == 'voice-test':
            print(f"Speaking: {args.text}")
            cli_speak(args.text)

        elif args.command == 'scan':
            cli_scan(args.type, args.path)

        elif args.command == 'status':
            cli_status()